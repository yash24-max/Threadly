"""Document parsing utilities."""

from pathlib import Path
from typing import Optional

from app.logger import get_logger

logger = get_logger(__name__)


def parse_document(file_path: str) -> str:
    """Parse document and extract text.

    Args:
        file_path: Path to document file

    Returns:
        Extracted text content

    Raises:
        ValueError: If file format not supported
    """
    path = Path(file_path)
    file_ext = path.suffix.lower()

    try:
        if file_ext == ".pdf":
            return _parse_pdf(file_path)
        elif file_ext == ".txt":
            return _parse_txt(file_path)
        elif file_ext == ".docx":
            return _parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    except Exception as e:
        logger.error(f"Error parsing document {file_path}: {e}")
        raise


def _parse_pdf(file_path: str) -> str:
    """Parse PDF file.

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text
    """
    try:
        from pypdf import PdfReader

        text = []
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {e}")
                    continue

        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        raise


def _parse_txt(file_path: str) -> str:
    """Parse plain text file.

    Args:
        file_path: Path to text file

    Returns:
        File content
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error parsing text file: {e}")
        raise


def _parse_docx(file_path: str) -> str:
    """Parse DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text
    """
    try:
        from docx import Document

        doc = Document(file_path)
        text = []

        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                text.append(" | ".join(row_cells))

        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        raise
